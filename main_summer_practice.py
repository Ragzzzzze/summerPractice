import sys
import json
from PyQt5.QtWidgets import (QApplication, QMainWindow, QTextBrowser, QAction, QFileDialog, QWidget, QMessageBox, QTextEdit,
                             QFontDialog, QColorDialog, QPushButton, QDialog, QComboBox, QLabel, QVBoxLayout, QInputDialog, 
                             QSpinBox, QGridLayout, QLineEdit)
from PyQt5 import uic, QtCore, QtGui
from PyQt5.QtGui import QTextCursor, QPixmap, QKeySequence, QFont, QTextCharFormat, QColor, QTextBlockFormat, QImage, QTextDocument, QBrush, QDesktopServices
from PyQt5.QtCore import QFileInfo, Qt, QUrl, QPoint, QRegExp, QEvent
from PyQt5.QtPrintSupport import QPrintDialog, QPrinter, QPrintPreviewDialog
from functools import partial
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import summer_practice.res_rc
import webbrowser
import fitz




class MarginsDialog(QDialog):
    DEFAULT_LINE_SPACING = 25

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Установка отступов и межстрочных интервалов")

        layout = QVBoxLayout()

        self.leftMargin = QSpinBox()
        self.leftMargin.setRange(0, 100)
        layout.addWidget(QLabel("Отступ слева:"))
        layout.addWidget(self.leftMargin)

        self.rightMargin = QSpinBox()
        self.rightMargin.setRange(0, 100)
        layout.addWidget(QLabel("Отступ справа:"))
        layout.addWidget(self.rightMargin)

        self.lineSpacing = QSpinBox()
        self.lineSpacing.setRange(0, 100)
        self.lineSpacing.setValue(0)  # Установка значения по умолчанию в 0
        layout.addWidget(QLabel("Межстрочный интервал:"))
        layout.addWidget(self.lineSpacing)

        self.applyButton = QPushButton("Применить")
        self.applyButton.clicked.connect(self.accept)
        layout.addWidget(self.applyButton)

        self.setLayout(layout)

    def getMargins(self):
        return {
            "left": self.leftMargin.value(),
            "right": self.rightMargin.value(),
            "lineSpacing": self.lineSpacing.value() or self.DEFAULT_LINE_SPACING
        }

    def setMargins(self, margins):
        self.leftMargin.setValue(margins.get("left", 0))
        self.rightMargin.setValue(margins.get("right", 0))
        self.lineSpacing.setValue(margins.get("lineSpacing", self.DEFAULT_LINE_SPACING) if margins.get("lineSpacing", self.DEFAULT_LINE_SPACING) != self.DEFAULT_LINE_SPACING else 0)


class StyleDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Выбор стиля")

        self.layout = QVBoxLayout(self)

        self.comboBox = QComboBox(self)
        self.layout.addWidget(QLabel("Выберите стиль"))
        self.layout.addWidget(self.comboBox)

        self.newStyleButton = QPushButton("Создать новый стиль", self)
        self.newStyleButton.clicked.connect(self.createNewStyle)
        self.layout.addWidget(self.newStyleButton)

        self.deleteStyleButton = QPushButton("Удалить стиль", self)
        self.deleteStyleButton.clicked.connect(self.deleteStyle)
        self.layout.addWidget(self.deleteStyleButton)

        self.applyButton = QPushButton("Применить", self)
        self.applyButton.clicked.connect(self.applyStyle)
        self.layout.addWidget(self.applyButton)

        self.styles = self.parent().loadStyles()  # Load styles from file
        self.updateComboBox()

    def updateComboBox(self):
        self.comboBox.clear()
        for style_name, style in self.styles.items():
            tooltip_text = (
                f"Шрифт: {style['font_family']}\n"
                f"Размер шрифта: {style['font_size']}\n"
                f"Цвет: {style['color']}\n"
                f"Выравнивание: {style['alignment']}\n"
                f"Отступ слева: {style['margins'].get('left', 0)}\n"
                f"Отступ справа: {style['margins'].get('right', 0)}\n"
                f"Межстрочный интервал: {style['lineSpacing']}"
            )
            self.comboBox.addItem(style_name)
            self.comboBox.setItemData(self.comboBox.count() - 1, tooltip_text, Qt.ToolTipRole)

    def createNewStyle(self):
        font, ok = QFontDialog.getFont()
        if ok:
            color = QColorDialog.getColor()
            if color.isValid():
                marginsDialog = MarginsDialog(self)
                marginsDialog.exec_()
                margins = marginsDialog.getMargins()

                styleName, ok = QInputDialog.getText(self, "Название стиля", "Введите название стиля:")
                if ok and styleName:
                    alignment, ok = QInputDialog.getItem(self, "Выравнивание", "Выберите выравнивание:",
                                                         ["Left", "Center", "Right", "Justify"], 0, False)
                    if ok:
                        lineSpacing = margins.pop("lineSpacing", MarginsDialog.DEFAULT_LINE_SPACING)
                        charFormat = QTextCharFormat()
                        charFormat.setFont(font)
                        charFormat.setForeground(color)

                        self.styles[styleName] = {
                            "font_family": font.family(),
                            "font_size": font.pointSize(),
                            "color": color.name(),
                            "alignment": alignment,
                            "margins": margins,
                            "lineSpacing": lineSpacing
                        }
                        self.updateComboBox()
                        self.parent().saveStyles(self.styles)
                        self.parent().applyTextStyle(charFormat, alignment, margins, lineSpacing)

    def applyStyle(self):
        try:
            styleName = self.comboBox.currentText()
            if styleName in self.styles:
                charFormat = QTextCharFormat()
                style = self.styles[styleName]
                charFormat.setFont(QFont(style["font_family"], style["font_size"]))
                charFormat.setForeground(QColor(style["color"]))

                alignment = style.get("alignment", "Left")
                margins = style.get("margins", {"left": 0, "right": 0})
                lineSpacing = style.get("lineSpacing", MarginsDialog.DEFAULT_LINE_SPACING)
                self.parent().applyTextStyle(charFormat, alignment, margins, lineSpacing)
            self.accept()
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Произошла ошибка при применении стиля: {str(e)}")

    def deleteStyle(self):
        styleName = self.comboBox.currentText()
        if styleName in self.styles:
            reply = QMessageBox.question(self, "Удаление стиля", f"Вы уверены, что хотите удалить стиль '{styleName}'?",
                                         QMessageBox.Yes | QMessageBox.No, QMessageBox.No)
            if reply == QMessageBox.Yes:
                del self.styles[styleName]
                self.updateComboBox()
                self.parent().saveStyles(self.styles)


class FindDialog(QDialog):
    def __init__(self, parent):
        super().__init__()
        self.parent = parent

        self.setWindowTitle("Find and Replace")

        layout = QGridLayout()
        self.finding_text = QLineEdit()
        layout.addWidget(QLabel("Find:"), 0, 0)
        layout.addWidget(self.finding_text, 0, 1)

        self.findButton = QPushButton("Find")
        layout.addWidget(self.findButton)

        self.comboBox = QComboBox()
        self.comboBox.addItems(["Select full word", "Select piece"])
        layout.addWidget(self.comboBox)

        self.setLayout(layout)

        self.extraSelections = []
        self.changed = False
        self.isFullWord = True

        self.setWhatsThis("Whats this")
        self.setWindowFlags(Qt.WindowContextHelpButtonHint | Qt.WindowCloseButtonHint)

        self.findButton.clicked.connect(self.find)

    def appendExtraSelection(self, tc):
            
            ex = QTextEdit.ExtraSelection()            
            ex.cursor = tc
            ex.format.setBackground(QBrush(Qt.yellow))
            self.extraSelections.append(ex)

    def find(self):
        self.comboBox.currentTextChanged.connect(self.onComboboxChanged)

        if self.changed: 
            self.extraSelections.clear()

        pattern = self.finding_text.text()

        cursor = self.parent.textEdit.textCursor()
        cursor.setPosition(0)
        doc = self.parent.textEdit.document()

        regex = QRegExp(pattern)

        self.isFullWord = True if self.comboBox.currentText() == "Select full word" else False

        pos = 0
        index = regex.indexIn(self.parent.textEdit.toPlainText(), pos)
        while (index != -1):

            cursor.setPosition(index)

            if self.isFullWord:
                cursor.movePosition(QTextCursor.EndOfWord, 1)
                if (cursor.selectionEnd() - cursor.selectionStart() == len(pattern)):
                    self.appendExtraSelection(cursor)


            if not self.isFullWord:
                cursor = doc.find(pattern, index)
                if not cursor.isNull(): 
                    self.appendExtraSelection(cursor)

            pos = index + regex.matchedLength()
            index = regex.indexIn(self.parent.textEdit.toPlainText(), pos)

        self.parent.textEdit.setExtraSelections(self.extraSelections)

    def closeEvent(self, event):
        self.extraSelections.clear()
        self.parent.textEdit.setExtraSelections(self.extraSelections)
        self.close()

    def onComboboxChanged(self, value):
        self.changed = True

        
class ReplaceDialog(QDialog):
    def __init__(self, parent):
        super().__init__()
        self.parent = parent

        self.setWindowTitle("Replace")

        self.gen_layout = QGridLayout()

        self.replacing_text = QLineEdit()
        self.gen_layout.addWidget(QLabel("Replace:"), 0, 0)
        self.gen_layout.addWidget(self.replacing_text, 0, 1)

        self.to_text = QLineEdit()
        self.gen_layout.addWidget(QLabel("To:"), 1, 0)
        self.gen_layout.addWidget(self.to_text, 1, 1)

        self.replaceButton = QPushButton("Replace")
        self.gen_layout.addWidget(self.replaceButton)

        self.setLayout(self.gen_layout) 

        self.replaceButton.clicked.connect(self.replace)


    def replace(self):
        self.replaceRec(0)


    def replaceRec(self, pos):
        
        doc = self.parent.textEdit.document()
        tc = QTextCursor(doc)

        pattern_for_replace = self.replacing_text.text()
        pattern_new = self.to_text.text()

        tc = doc.find(pattern_for_replace, pos)
        if not tc.isNull():
            tc.removeSelectedText()
            tc.insertText(pattern_new)
            self.replaceRec(tc.selectionEnd())


class HrefDialog(QDialog):
    def __init__(self, parent):
        super().__init__(parent)

        self.parent = parent

        self.setWindowTitle("Insert hyperlink")


        layout = QGridLayout()
        self.adress = QLineEdit()
        layout.addWidget(QLabel("Adress:"), 0, 0)
        layout.addWidget(self.adress, 0, 1)

        self.word = QLineEdit()
        layout.addWidget(QLabel("Text:"), 1, 0)
        layout.addWidget(self.word, 1, 1)

        self.insertButton = QPushButton("Insert")
        layout.addWidget(self.insertButton)

        self.setLayout(layout)

        self.insertButton.clicked.connect(self.insertHyperlink)

    def insertHyperlink(self):
        cursor = self.parent.textEdit.textCursor()

        format = cursor.charFormat()
        format_original = cursor.charFormat()
        format.setForeground(QColor("blue"))
        format.setFontUnderline(True)

        if not self.adress.text():
            QMessageBox.critical(self, "Error", "Adress line is empty")
        
        if not self.word.text():
            self.word.setText(self.adress.text())
        
        format.setAnchor(True)
        format.setAnchorHref(self.adress.text())
        format.setToolTip(self.adress.text())

        cursor.insertText(self.word.text(), format)
        cursor.insertText(" ", format_original)

        self.close()


class MyWidget(QMainWindow):
    def __init__(self):
        super().__init__()
        # Загрузка UI из файла
        uic.loadUi('main_practice_summer.ui', self)

        initial_font = QFont("Calibri", 15)
        self.textEdit.setFont(initial_font)
        self.textEdit.setTabStopDistance(20)

        self.pages = QSpinBox()
        self.toolBar.addWidget(self.pages)

        self.pages.valueChanged.connect(self.change_page)
        self.page_contents = {}  # Хранение содержимого страниц
        self.current_page = 1
        self.pages.setMinimum(1)
        self.pages.setValue(1)
        self.load_page_content()

        self.actionCreate_Style.triggered.connect(self.openStyleDialog)

        self.actionImage.triggered.connect(self.insertImageAction)

        self.actionNew.triggered.connect(self.fileNew)
        self.actionNew_3.triggered.connect(self.fileNew)
        self.actionOpen_2.triggered.connect(self.openFile)
        self.actionOpen.triggered.connect(self.openFile)
        self.actionSave.triggered.connect(self.save_as_pdf)
        self.actionSave_2.triggered.connect(self.save_as_pdf)
        self.actionPrint_3.triggered.connect(self.printfile)
        self.actionPrint_2.triggered.connect(self.printfile)
        self.actionPrint_Preview.triggered.connect(self.printPreview)
        self.actionPrint_Preview_2.triggered.connect(self.printPreview)
        self.actionExit_3.triggered.connect(self.exitEditor)
        self.actionCopy_2.triggered.connect(self.copy)
        self.actionCopy.triggered.connect(self.copy)
        self.actionPaste_2.triggered.connect(self.paste)
        self.actionPaste.triggered.connect(self.paste)
        self.actionCut.triggered.connect(self.cut)
        self.actionCut_2.triggered.connect(self.cut)
        self.actionUndo.triggered.connect(self.textEdit.undo)
        self.actionUndo_2.triggered.connect(self.textEdit.undo)
        self.actionRedo.triggered.connect(self.textEdit.redo)
        self.actionRedo_2.triggered.connect(self.textEdit.redo)
        self.actionFont.triggered.connect(self.fontDialog)
        self.actionFont_2.triggered.connect(self.fontDialog)
        self.actionColor.triggered.connect(self.colorDialog)
        self.actionLeft_2.triggered.connect(self.setAlignment)
        self.actionCenter_2.triggered.connect(self.setAlignment)
        self.actionRight_2.triggered.connect(self.setAlignment)
        self.actionJustify.triggered.connect(self.setAlignment)
        self.actionMargins.triggered.connect(self.openMarginsDialog)
        self.actionMargins_2.triggered.connect(self.openMarginsDialog)

        self.actionFind.triggered.connect(self.findWindow)
        self.actionReplace.triggered.connect(self.replaceWindow)
        self.actionInsert_new.triggered.connect(self.insert)

        self.actionBold.setShortcut(QKeySequence.Bold)
        self.actionBold.setCheckable(True)
        self.actionBold.toggled.connect(lambda x: self.textEdit.setFontWeight(QFont.Bold if x else QFont.Normal))

        self.actionItalic.setShortcut(QKeySequence.Italic)
        self.actionItalic.setCheckable(True)
        self.actionItalic.toggled.connect(self.textEdit.setFontItalic)

        self.actionUnderline.setShortcut(QKeySequence.Underline)
        self.actionUnderline.setCheckable(True)
        self.actionUnderline.toggled.connect(self.textEdit.setFontUnderline)
        self.setInitialFont()

        self.actionLeft_2.setShortcut("CTRL+L")
        self.actionLeft_2.setCheckable(True)
        self.actionLeft_2.triggered.connect(partial(self.setAlignment, Qt.AlignLeft))

        self.actionCenter_2.setShortcut("CTRL+E")
        self.actionCenter_2.setCheckable(True)
        self.actionCenter_2.triggered.connect(partial(self.setAlignment, Qt.AlignCenter))

        self.actionRight_2.setShortcut("CTRL+R")
        self.actionRight_2.setCheckable(True)
        self.actionRight_2.triggered.connect(partial(self.setAlignment, Qt.AlignRight))

        self.actionJustify.setShortcut("CTRL+J")
        self.actionJustify.setCheckable(True)
        self.actionJustify.triggered.connect(partial(self.setAlignment, Qt.AlignJustify))

        # Load styles from file at startup
        self.styles = self.loadStyles()

        self.currentMargins = {"left": 0, "right": 0}
        self.actionDefault.triggered.connect(self.applyDefaultStyle)

    def applyDefaultStyle(self):
        styleName = "Default"
        if styleName in self.styles:
            charFormat = QTextCharFormat()
            style = self.styles[styleName]
            charFormat.setFont(QFont(style["font_family"], style["font_size"]))
            charFormat.setForeground(QColor(style["color"]))

            alignment = style.get("alignment", "Left")
            margins = style.get("margins", {"left": 0, "right": 0})
            lineSpacing = style.get("lineSpacing", MarginsDialog.DEFAULT_LINE_SPACING)

            self.applyTextStyle(charFormat, alignment, margins, lineSpacing)

    def setAlignment(self, alignment):
        try:
            self.textEdit.setAlignment(alignment)

            if alignment == Qt.AlignLeft:
                self.actionCenter_2.setChecked(False)
                self.actionRight_2.setChecked(False)
                self.actionJustify.setChecked(False)

            elif alignment == Qt.AlignRight:
                self.actionLeft_2.setChecked(False)
                self.actionCenter_2.setChecked(False)
                self.actionJustify.setChecked(False)

            elif alignment == Qt.AlignCenter:
                self.actionRight_2.setChecked(False)
                self.actionLeft_2.setChecked(False)
                self.actionJustify.setChecked(False)

            elif alignment == Qt.AlignJustify:
                self.actionRight_2.setChecked(False)
                self.actionLeft_2.setChecked(False)
                self.actionCenter_2.setChecked(False)
        except Exception as E:
            pass

    def setInitialFont(self):
        initial_font = QFont("Calibri", 14)
        self.textEdit.setFont(initial_font)

    def insertImageAction(self):
        file_name, _ = QFileDialog.getOpenFileName(self, 'Выбрать изображение', '',
                                                   'Изображения (*.png *.jpg *.bmp *.gif)')
        if file_name:
            cursor = self.textEdit.textCursor()
            cursor.insertHtml('<img src="{}">'.format(file_name))

    def fileNew(self):
        result = QMessageBox.question(
            self,
            "Открытие нового файла",
            "Хотите сохранить изменения в файл?",
            QMessageBox.Save | QMessageBox.Discard | QMessageBox.Cancel
        )

        if result == QMessageBox.Save:
            self.saveFile()

        elif result == QMessageBox.Discard:
            self.textEdit.clear()

        elif result == QMessageBox.Cancel:
            return

    def save_as_pdf(self):
        # Сначала сохраняем текущую страницу
        self.page_contents[self.current_page] = self.textEdit.toHtml()

        file_path, _ = QFileDialog.getSaveFileName(self, "Save File", "", "PDF Files (*.pdf)")
        if file_path:
            try:
                printer = QPrinter(QPrinter.HighResolution)
                printer.setOutputFormat(QPrinter.PdfFormat)
                printer.setOutputFileName(file_path)

                # Создаем новый документ для финального PDF
                final_document = QTextDocument()

                # Объединяем содержимое всех страниц в один HTML документ
                full_html = "<html><body>"
                for page in sorted(self.page_contents.keys()):
                    text = self.page_contents.get(page, "")
                    if page != min(self.page_contents.keys()):  # Добавляем разрыв страницы между страницами
                        full_html += "<div style='page-break-before:always;'></div>"
                    full_html += f"<div>{text}</div>"

                full_html += "</body></html>"

                final_document.setHtml(full_html)
                final_document.print_(printer)
                QMessageBox.information(self, "Сохранение завершено", f"Документ сохранен по пути {file_path}")

            except Exception as e:
                QMessageBox.critical(self, "Ошибка сохранения", f"Произошла ошибка при сохранении PDF: {str(e)}")

    def change_page(self):
        # Сохраняем содержимое и формат текущей страницы
        self.page_contents[self.current_page] = self.textEdit.toHtml()

        # Меняем страницу
        self.current_page = self.pages.value()
        self.load_page_content()

    def load_page_content(self):
        # Загружаем содержимое и формат для текущей страницы
        text = self.page_contents.get(self.current_page, "")
        self.textEdit.setHtml(text)

    def openFile(self):
        filePath, _ = QFileDialog.getOpenFileName(self, 'Open File', '',
                                                  'Text Files (*.txt);;Word Documents (*.docx);;PDF Files (*.pdf);;All Files (*)')
        if filePath:
            if filePath.endswith('.txt'):
                with open(filePath, 'r') as file:
                    data = file.read()
                    self.textEdit.setPlainText(data)
            elif filePath.endswith('.docx'):
                self.openDocxFile(filePath)
            elif filePath.endswith('.pdf'):
                self.openPdfFile(filePath)

    def openDocxFile(self, filePath):
        document = Document(filePath)
        text = '\n'.join([paragraph.text for paragraph in document.paragraphs])
        self.textEdit.setPlainText(text)

    def openPdfFile(self, filePath):
        try:
            doc = fitz.open(filePath)
            text = ""
            for page_num in range(doc.page_count):
                page = doc.load_page(page_num)
                text += page.get_text("text")
            self.textEdit.setPlainText(text)
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось открыть файл PDF: {str(e)}")

    def printfile(self):
        printer = QPrinter(QPrinter.HighResolution)
        dialog = QPrintDialog(printer, self)

        if dialog.exec_() == QPrintDialog.Accepted:
            self.textEdit.print_(printer)

    def printPreview(self):
        printer = QPrinter(QPrinter.HighResolution)
        previewDialog = QPrintPreviewDialog(printer, self)
        previewDialog.paintRequested.connect(self.printDocument)
        previewDialog.exec_()

    def printDocument(self, printer):
        # Используйте QTextDocument для печати
        document = QTextDocument()

        # Объединяем содержимое всех страниц в один HTML документ
        full_html = "<html><body>"
        for page in sorted(self.page_contents.keys()):
            text = self.page_contents.get(page, "")
            if page != min(self.page_contents.keys()):  # Добавляем разрыв страницы между страницами
                full_html += "<div style='page-break-before:always;'></div>"
            full_html += f"<div>{text}</div>"
        full_html += "</body></html>"

        document.setHtml(full_html)

        # Печать всего документа
        document.print_(printer)

    def exitEditor(self):
        result = QMessageBox.question(
            self,
            "Exit",
            "Do you really want to exit?",
            QMessageBox.No | QMessageBox.Yes
        )

        if result == QMessageBox.Yes:
            self.close()

    def copy(self):
        cursor = self.textEdit.textCursor()
        textSelected = cursor.selectedText()
        clipboard = QApplication.clipboard()
        clipboard.setText(textSelected)

    def paste(self):
        clipboard = QApplication.clipboard()
        textToPaste = clipboard.text()
        self.textEdit.insertPlainText(textToPaste)

    def cut(self):
        cursor = self.textEdit.textCursor()
        textSelected = cursor.selectedText()
        self.copiedText = textSelected
        self.textEdit.cut()

    def fontDialog(self):
        cursor = self.textEdit.textCursor()
        if cursor.hasSelection():
            font, ok = QFontDialog.getFont()
            if ok:
                format = QTextCharFormat()
                format.setFont(font)
                cursor.mergeCharFormat(format)

    def colorDialog(self):
        cursor = self.textEdit.textCursor()
        if cursor.hasSelection():
            color = QColorDialog.getColor()
            if color.isValid():
                format = QTextCharFormat()
                format.setForeground(color)
                cursor.mergeCharFormat(format)

    def openStyleDialog(self):
        dialog = StyleDialog(self)
        dialog.exec_()

    def applyTextStyle(self, charFormat, alignment, margins=None, lineSpacing=None):
        cursor = self.textEdit.textCursor()
        blockFormat = cursor.blockFormat()
        blockFormat.setAlignment(self.getAlignmentFromString(alignment))

        if margins:
            blockFormat.setLeftMargin(margins.get("left", 0))
            blockFormat.setRightMargin(margins.get("right", 0))

        if lineSpacing is not None:
            blockFormat.setLineHeight(lineSpacing, QTextBlockFormat.FixedHeight)

        if cursor.hasSelection():
            cursor.mergeCharFormat(charFormat)
            cursor.mergeBlockFormat(blockFormat)
        else:
            cursor.setBlockFormat(blockFormat)
            cursor.setCharFormat(charFormat)
            self.textEdit.setTextCursor(cursor)

        self.textEdit.setFocus()

    def getAlignmentFromString(self, alignment_string):
        if alignment_string == "Left":
            return Qt.AlignLeft
        elif alignment_string == "Center":
            return Qt.AlignCenter
        elif alignment_string == "Right":
            return Qt.AlignRight
        elif alignment_string == "Justify":
            return Qt.AlignJustify
        return Qt.AlignLeft  # Значение по умолчанию

    def loadStyles(self):
        try:
            with open('styles.json', 'r') as file:
                data = file.read()
                # Check if the file is empty
                if not data:
                    return {}
                return json.loads(data)
        except (FileNotFoundError, json.JSONDecodeError):
            return {}

    def saveStyles(self, styles):
        try:
            with open('styles.json', 'w') as file:
                json.dump(styles, file, indent=4)
        except IOError as e:
            QMessageBox.critical(self, "Error", f"Failed to save styles: {str(e)}")

    def keyPressEvent(self, event):
        if event.key() in (Qt.Key_Return, Qt.Key_Enter):
            cursor = self.textEdit.textCursor()
            cursor.movePosition(QTextCursor.StartOfLine, QTextCursor.MoveAnchor)
            cursor.movePosition(QTextCursor.EndOfLine, QTextCursor.KeepAnchor)
            cursor.removeSelectedText()
            cursor.deleteChar()

            default_format = QTextCharFormat()
            default_format.setFont(QFont("Calibri", 11))
            self.textEdit.setCurrentCharFormat(default_format)

            super().keyPressEvent(event)

            # Добавляем пустую строку
            cursor.insertBlock()
        else:
            super().keyPressEvent(event)

    def mousePressEvent(self, event):
        x, y = event.pos().x(), event.pos().y()
        new_pos = QPoint(x, y - 61)
        self.anchor = self.textEdit.anchorAt(new_pos)
        
        if self.anchor:
            QApplication.setOverrideCursor(Qt.PointingHandCursor)

            QDesktopServices.openUrl(QUrl(self.anchor))
            QApplication.setOverrideCursor(Qt.ArrowCursor)
            self.anchor = None
 
    def openMarginsDialog(self):
        dialog = MarginsDialog(self)
        dialog.setMargins(self.currentMargins)  # Set current margins
        if dialog.exec_():
            margins = dialog.getMargins()
            self.applyMargins(margins)
            self.currentMargins = margins  # Update current margins

    def applyMargins(self, margins):
        cursor = self.textEdit.textCursor()
        blockFormat = cursor.blockFormat()

        blockFormat.setLeftMargin(margins["left"])
        blockFormat.setRightMargin(margins["right"])

        # Устанавливаем минимальное значение для межстрочного интервала
        lineSpacing = max(margins.get("lineSpacing", 0), 1)
        blockFormat.setLineHeight(lineSpacing, QTextBlockFormat.FixedHeight)

        if cursor.hasSelection():
            cursor.mergeBlockFormat(blockFormat)
        else:
            cursor.setBlockFormat(blockFormat)
            self.textEdit.setTextCursor(cursor)

    def findWindow(self):
        self.dialog = FindDialog(self)
        self.dialog.exec_()
    
    def replaceWindow(self):
        self.dialog = ReplaceDialog(self)
        self.dialog.exec_()

    def insert(self):
        self.dialog = HrefDialog(self)
        self.dialog.exec_()


if __name__ == '__main__':
    app = QApplication(sys.argv)
    window = MyWidget()
    window.show()
    sys.exit(app.exec_())
